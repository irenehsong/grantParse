"""
This program is free software: you can redistribute it and/or modify
it under the terms of the GNU General Public License as published by
the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
GNU General Public License for more details.

You should have received a copy of the GNU General Public License
along with this program.  If not, see <https://www.gnu.org/licenses/>.
"""

import datetime
import html as html
import os
import threading
import xml.etree.ElementTree as et
from tkinter import (BOTTOM, LEFT, RIGHT, TOP, Button, Frame, Label, Entry, Tk,
                     messagebox)

import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from tkcalendar import DateEntry

import GrantDownloader
import word

# dictionary of agencies using agency code as key
# these were all the agencies in the search function for Grants.gov
# I added 'N/A' to the list to make sure that if we did not have a match, we would still have a key for it
agencyDictionary = {'USAID': 'Agency for International Development',
                    'AC': 'AmeriCorps',
                    'USDA': 'Department of Agriculture',
                    'DOC': 'Department of Commerce',
                    'DOE': 'Department of Energy',
                    'DOD': 'Department of Defense',
                    'ED': 'Department of Education',
                    'PAMS': 'Department of Health and Human Services',
                    'HHS': 'Department of Health and Human Services',
                    'DHS': 'Department of Homeland Security',
                    'HUD': 'Department of Housing and Urban Development',
                    'USDOJ': 'Department of Justice',
                    'DOL': 'Department of Labor',
                    'DOS': 'Department of State',
                    'DOI': 'Department of the Interior',
                    'USDOT': 'Department of the Treasury',
                    'DOT': 'Department of Transportation',
                    'VA': 'Department of Veterans Affairs',
                    'EPA': 'Environmental Protection Agency',
                    'GCERC': 'Gulf Coast Ecosystem Restoration Council',
                    'IMLS': 'Institute of Museum and Library Services',
                    'MCC': 'Millennium Challenge Corportation',
                    'NASA': 'National Aeronautics and Space Administration',
                    'NARA': 'National Archives and Records Administration',
                    'NEA': 'National Endowment for the Arts',
                    'NEH': 'National Endowment for the Humanities',
                    'NSF': 'National Science Foundation',
                    'NRC': 'National Resource Conservation Council',
                    'SBA': 'Small Business Administration',
                    'SSA': 'Social Security Administration',
                    'N/A': 'Other Agencies'}

# String that saves the http portion of the XML tags so that we don't have to keep typing it out,
# e.g. referencing the tag '{http://apply.grants.gov/system/OpportunityDetail-V1.0}OpportunityTitle'
# becomes just linkString+'OpportunityTitle'
linkString = '{http://apply.grants.gov/system/OpportunityDetail-V1.0}'

# ********************************************DEF*****************************************************************

# Convert our dates into a better looking format with slashes, MM/DD/YYYY


def dateConversion(date):
    newDate = date[:2] + "/" + date[2:4] + "/" + date[4:]
    return newDate

# Convert our dates to format similar to January 01, 2021


def dateStringVersion(date):
    newDate = ''
    if (date != 'N/A'):

        str(date)
        monthList = ['January ', 'February ', 'March ', 'April ', 'May ', 'June ',
                     'July ', 'August ', 'September ', 'October ', 'November ', 'December ']
        monthNum = int(date[:2])
        temp = monthList[monthNum - 1]
        newDate = temp + date[2:4] + ", " + date[4:]
    else:
        newDate = 'N/A'
    return newDate


# input a string to add commas
# example input : 10000000
# example output: 10,000,000
def addCommasAndDollarSign(amountStr):
    # check if string is a number, if not return the string back
    if not amountStr.isnumeric():
        return amountStr
    else:
        return "${:,}".format(int(amountStr))


# convert our dates into a year, month, day hierarchy so that earlier dates are natrually smaller numbers (strings in this case) than later dates
def dateHierarchyForm(date):
    newDate = date[4:] + date[:4]
    return newDate


# this function takes an agency code and returns the name of the agency
def generateAgencyName(agencyCode):
    # agency codes have dashes to separate the information.
    if '-' in agencyCode:
        # we only want the first part of the agency code which identifies the agency itself
        agencyCode = agencyCode.split('-')[0]

    # if the agency is in the list of agencies, we can use the agency code to get the name
    if agencyCode in agencyDictionary:
        agencyCode = agencyDictionary[agencyCode]   # of the agency

    else:
        # if the agency is not in the list of agencies, we will use the string 'Other Agencies'
        agencyCode = 'Other Agencies'

    return agencyCode

# function to generate link to grant from grant ID


def generateLink(grantID):

    link = (
        f"https://www.grants.gov/web/grants/view-opportunity.html?oppId={grantID}")
    return link

# function to reduce the number of words for a string
# will help to reduce the number of words in the description


def wordLimiter(string, limit):
    string = string.split()[:limit]
    string = " ".join(string) + "..."
    return string

# create table of unique agencies (want to improve this, lot of redudancies related
# to multiple instances of the same agency with additional information)


def tableOfContents(listA, agency):

    if agency in listA:                     # if the agency is already in the list, we don't need to add it again
        return listA
    else:
        # if the agency is not in the list, we add it to the list
        listA.append(agency)

    return listA


# function to return an attribute of interest from a given opportunity
def getOpportunityInfo(opportunity, attribute):
    #
    # Store each text of an attribute as a string, or store as 'N/A' if none exist
    myInfo = getattr(opportunity.find(linkString + attribute), 'text', 'N/A')
    # at given attribute location
    return myInfo


# ********************************************MAIN Object*****************************************************************
class Grant:

    def __init__(self, agencyCode, agencyName, opportunityTitle, postDate, dueDate, numAwards,
                 totalFunding, awardCeiling, awardFloor, oppNumber, description, grantLink, contactInfo, eligApplicants='N/A'):

        self.agencyCode = agencyCode
        self.distinctAgency = generateAgencyName(agencyCode)
        self.agencyName = agencyName
        self.opportunityTitle = opportunityTitle
        self.postDate = postDate
        self.dueDate = dueDate
        self.numAwards = numAwards
        self.totalFunding = addCommasAndDollarSign(totalFunding)
        self.awardCeiling = addCommasAndDollarSign(awardCeiling)
        self.awardFloor = addCommasAndDollarSign(awardFloor)
        self.oppNumber = oppNumber
        self.description = description
        self.eligApplicants = eligApplicants
        self.grantLink = grantLink
        self.contactInfo = contactInfo

# ********************************************MAIN FUNCTIONS*****************************************************************


def printGrant(grant):
    print("Agency name:                     " + grant.agencyName)
    print("Opportunity title:               " + grant.opportunityTitle)
    print("Post date:                       " +
          grant.postDate)
    print("Due date:                        " +
          grant.dueDate)
    print("Expected Number of awards:       " + grant.numAwards)
    print("Estimated total program funding: " +
          addCommasAndDollarSign(grant.totalFunding))
    print("Award Ceiling:                   " +
          addCommasAndDollarSign(grant.awardCeiling))
    print("Award floor:                     " +
          addCommasAndDollarSign(grant.awardFloor))
    print("Funding opportunity number:      " + grant.oppNumber)
    print()
    print("Purpose: " + grant.description)
    print()
    print("Eligible applicants: " + grant.eligApplicants)
    print()
    print("Contact information: " + grant.contactInfo)
    print()
    print("Link: " + grant.grantLink)

    print()


# function to create a dictionary using the distinctAgency as a key and the grants as values
def grantDictionaryAdd(grantDictionary, grant):
    grantDictionary.setdefault(grant.distinctAgency, []).append(grant)
    return grantDictionary

# ********************************DRIVER_CODE****************************************************************************

# --------------------------- UI BEGIN ---------------------------

# Set today's date and set last week's date
today = datetime.date.today()
last_week = today - datetime.timedelta(days=7)

# Basic Settings, window title / size
root = Tk()
root.title('US Government Grant Report Tool')
if os.name == 'posix':
    root.iconbitmap('@resource/tux.xbm')
else:
    root.iconbitmap('resource/icon.ico')
root.geometry("500x320")

# Sets layout of modules
top = Frame(root)
bottom = Frame(root, width=100)
top.pack(side=TOP)
bottom.pack(side=BOTTOM, fill=None, expand=False)

# Sets padding and text of UI Label
my_toplabel = Label(root, text="Please select a date range.")
my_toplabel.pack(pady=10, in_=top)

# Set and post date entry fields (first set to 1 week past, second to current date)
calone = DateEntry(root, width=12, background='darkblue',
                   foreground='white', borderwidth=2, year=last_week.year, month=last_week.month, day=last_week.day)
calone.pack(in_=top, side=LEFT, padx=20, pady=10)

caltwo = DateEntry(root, width=12, background='darkblue',
                   foreground='white', borderwidth=2)
caltwo.pack(in_=top, side=RIGHT, padx=20, pady=10)

# Set URL entry field
grant_url_label = Label(root, text="Please enter the URL of the grant information web.")
grant_url_label.pack(pady=10, side=TOP)
grant_url = Entry(root, width=150)
grant_url.insert(0, "https://www.grants.gov/xml-extract")
grant_url.pack(side=TOP, padx=20, pady=10)

# Function for collecting date from calendar and url
def grab_data():
    global userdateone, userdatetwo, userurl
    userdateone = calone.get_date()
    userdatetwo = caltwo.get_date()
    userurl = grant_url.get()

    if userdateone > userdatetwo or userurl == "":
        if userdateone > userdatetwo:
            messagebox.showerror(
            "Improper Date Range", "Please ensure your first date is before your second date.")
        if userurl == "":
            messagebox.showerror(
                "Enter the URL.", "Please enter the URL.")
    else:
        root.destroy()

# Function for usage in multithreading of grantdownloader.py
def downloadxml():
    global mytree, userurl
    mytree = et.parse(GrantDownloader.get(userurl))

userdateone = calone.get_date()
userdatetwo = caltwo.get_date()

# Button Grabs selected dates then closes if userdateone > userdatetwo
my_button = Button(root, text="Confirm",
                   activebackground='gray', command=grab_data)
my_button.pack(pady=10, padx=10, in_=bottom, side=RIGHT)

# Location for date to post
my_label = Label(root, text="")
my_label.pack(pady=10)

# loopy boi
root.mainloop()

# Convert datetime object to string for comparison
dateRangeOne = userdateone.strftime("%Y%m%d")
dateRangeTwo = userdatetwo.strftime("%Y%m%d")

'''
Trying to look into more efficient parsing, this site seems to have a better, space efficient parsing method I need to explore:
https://newbedev.com/efficient-way-to-iterate-through-xml-elements
'''

# Parse our file and store it in a tree
mytree = et.parse(GrantDownloader.get(userurl))

# Set what function other thread will execute
th = threading.Thread(target=downloadxml)

# Begin multithread download on UI load
th.start()

# Print message telling user that the document is being generated
print("Generating grants report...")

# --------------------------- UI END ---------------------------

# Waits for both threads to finish their execution before continuing
th.join()

###############################################################---XML Parsing/Grant Generation---#############################################################################

# Store the root element of this file
myroot = mytree.getroot()

# Count the number of grants that we have chosen to print,
count = 0

# Declare our list to store agency names
agencyList = []
grantDictionary = {}

# Check each grant opportunity in our tree
for opportunity in myroot:

    # Get the postdate first
    # getattr(opportunity.find(linkString + 'PostDate'), 'text', 'N/A')
    postDate = getOpportunityInfo(opportunity, 'PostDate')
    # Set the desired earliest date
    # Finds grants of date range selected in UI
    if dateRangeOne <= dateHierarchyForm(postDate) <= dateRangeTwo:

        # print('************************************************************************************************************************')
        # print()

        # Store each text of qualifying grants as a string, or store as 'N/A' if none exist
        grant = Grant(
            agencyCode=getOpportunityInfo(opportunity, 'AgencyCode'),
            agencyName=getOpportunityInfo(opportunity, 'AgencyName'),
            opportunityTitle=html.unescape(
                getOpportunityInfo(opportunity, 'OpportunityTitle')),
            postDate=dateStringVersion(
                getOpportunityInfo(opportunity, 'PostDate')),
            dueDate=dateStringVersion(
                getOpportunityInfo(opportunity, 'CloseDate')),
            numAwards=getOpportunityInfo(
                opportunity, 'ExpectedNumberOfAwards'),
            totalFunding=getOpportunityInfo(
                opportunity, 'EstimatedTotalProgramFunding'),
            awardCeiling=getOpportunityInfo(opportunity, 'AwardCeiling'),
            awardFloor=getOpportunityInfo(opportunity, 'AwardFloor'),
            oppNumber=getOpportunityInfo(opportunity, 'OpportunityNumber'),
            description=html.unescape(
                getOpportunityInfo(opportunity, 'Description')),
            eligApplicants=html.unescape(getOpportunityInfo(
                opportunity, 'AdditionalInformationOnEligibility')),
            grantLink=generateLink(getOpportunityInfo(
                opportunity, 'OpportunityID')),
            contactInfo=html.unescape(getOpportunityInfo(
                opportunity, 'GrantorContactText'))
        )
        tableOfContents(agencyList, grant.distinctAgency)
        # Create a grant object
        grantDictionary = grantDictionaryAdd(grantDictionary, grant)
        # printGrant(grant)

        # Count the selected grant
        count += 1

# print('**********************************************************************************************************')
# # Print out the number of grants that qualified. I used this to check to make sure pruning was happening
# print("Number of grants", count)

# # sort the list of agencies
# print()
agencyList.sort()

# print the list of agencies
# print('Table of Contents')
# print('----------------------------------------------------------------------')
# for x in agencyList:
#     print(x)

#####################################################################################################################
#  ChangeTemplate

#! Opens the word templet file. Change the string passed as an argument to the method .Document to use a different template
# Provided templates are:
#   Marshall template.docx
#   OpsWatch template.docx

doc = docx.Document("Marshall template.docx")

#####################################################################################################################

#! change text in paragraph 9 to the number of grants
doc.paragraphs[9].text = str(datetime.date.today().strftime("%B %d, %Y"))

#! Random paragraph object to position the start of the hyperlink prints
spacerpara = doc.add_paragraph("\n")
paragraph_format = spacerpara.paragraph_format
paragraph_format.line_spacing = 1.0

#! Add page break
doc.add_page_break()

#! Add Header to start of Grants sections

line = doc.add_paragraph()
line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = line.add_run("\nGrants\n")
run.bold = True
font = run.font
font.size = Pt(22)
font.name = 'Times New Roman'
font.underline = True

#! Set pointer to equal paragraph 11
paracount = 12
pointer = doc.paragraphs[paracount]

#! Check if it's a new agency
agency_check = set()
agency_name_check = set()

#! This prints generates the bookmarks
for index, agency in enumerate(agencyList):

    grantDictionary[agency].sort(key=lambda x: x.dueDate)

    #! Loop over each grant in the dictionary
    grant_list = grantDictionary.get(agency)
    for i in grant_list:

        if agency not in agency_check:
            agency_check.add(agency)
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = 1.0
            bookmark_para = word.add_bookmark(
                paragraph, agency, f"bookmark{str(index)}")

            paragraph_format = pointer.paragraph_format
            paragraph_format.line_spacing = 1.0
            word.insert_paragraph_after(pointer, word.add_link(
                pointer, f"bookmark{str(index)}", agency))
            paracount += 1
            pointer = doc.paragraphs[paracount]

        if i.agencyName not in agency_name_check:
            agency_name_check.add(i.agencyName)

            paragraph_format = pointer.paragraph_format
            paragraph_format.line_spacing = 1.0
            word.insert_paragraph_after(pointer, i.agencyName)
            paracount += 1
            pointer = doc.paragraphs[paracount]

            paragraph_format = pointer.paragraph_format
            paragraph_format.line_spacing = 1.0
            word.insert_paragraph_after(pointer, f"\t• {i.opportunityTitle}")
            paracount += 1
            pointer = doc.paragraphs[paracount]

        else:
            paragraph_format = pointer.paragraph_format
            paragraph_format.line_spacing = 1.0
            word.insert_paragraph_after(pointer, f"\t• {i.opportunityTitle}")
            paracount += 1
            pointer = doc.paragraphs[paracount]

        paragraph = doc.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 1.0

        paragraph.add_run(f"\nAgency Name: {i.agencyName}").bold = True
        paragraph.add_run(
            f"\nOpportunity Title: {i.opportunityTitle}").bold = True
        paragraph.add_run(f"\nPost Date:\t\t\t\t\t\t{i.postDate}").bold = True
        paragraph.add_run(
            f"\nProposal Due Date:\t\t\t\t\t{i.dueDate}").bold = True
        paragraph.add_run(
            f"\nExpected Number of awards:\t\t\t{i.numAwards}").bold = True
        paragraph.add_run(
            f"\nEstimated total program funding:\t\t{i.totalFunding}").bold = True
        paragraph.add_run(
            f"\nAward Ceiling:\t\t\t\t\t{i.awardCeiling}").bold = True
        paragraph.add_run(
            f"\nAward Floor:\t\t\t\t\t{i.awardFloor}").bold = True
        paragraph.add_run(
            f"\nFunding Opportunity Number:\t\t\t{i.oppNumber}").bold = True

        run = paragraph.add_run(f"\n\nPurpose: ")
        run.bold = True

        run = paragraph.add_run(f"{i.description}")
        font = run.font
        font.size = Pt(12)
        # font.italic = True
        # font.name = 'Times New Roman'

        # Print eligibility information
        run = paragraph.add_run(f"\n\nEligible Applicants: ")
        run.bold = True

        run = paragraph.add_run(f"{i.eligApplicants}")
        font = run.font
        font.size = Pt(12)
        #font.name = 'Times New Roman'
        paragraph.add_run(f"\n")

        # Print contact information if available
        if (i.contactInfo != 'N/A'):
            run = paragraph.add_run(f"\nContact: ")
            run.bold = True

            contactArr = i.contactInfo.split('<br/>')
            for j in contactArr:
                run = paragraph.add_run(f"\n{j}")

            paragraph.add_run(f"\n")

        #! Add hyperlink to grant
        link_para = doc.add_paragraph()
        word.add_hyperlink(link_para, f"{i.grantLink}\n", i.grantLink)

    else:
        #! Random paragraph object to position the start of the next agency name better
        paragraph_format.line_spacing = 1.0
        word.insert_paragraph_after(pointer, "\n")
        paracount += 1
        pointer = doc.paragraphs[paracount]
        doc.add_page_break()

doc.save(f"GrantsReport_{today}.docx")
